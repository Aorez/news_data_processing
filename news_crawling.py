import random
import re
import time
import requests
import pymongo
from lxml import etree
import csv
# 导入模块
from docx import Document
# 此模块中包含 docx 中各类单位方法
from docx import shared

# 连接MongoDB数据库
mongodb_client = pymongo.MongoClient('mongodb://localhost:27017/')
dblist = mongodb_client.list_database_names()
if "db_news" in dblist:
    # print("数据库已存在！")
    eval("1+1")
db_news = mongodb_client['db_news']
collection_list = db_news.list_collection_names()
if "collection_news" in collection_list:  # 判断 sites 集合是否存在
    # print("集合已存在！")
    eval("1+1")
collection_news = db_news['collection_news']
collection_news.delete_many({})

# 待爬网址
base_url = "https://news.china.com/domestic/index.html"
# 反反爬
headers = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/102.0.0.0 "
                  "Safari/537.36",
    "Accept-Encoding": "gzip, deflate, br",
    "Accept-Language": "zh-CN,zh;q=0.9"
}

# 爬取信息
base_response = requests.get(url=base_url, headers=headers)
# 设置编码
base_response.encoding = "utf-8"
# 把信息通过etree进行赋值
base_html = etree.HTML(base_response.text)
# xpath定位新闻列表信息
news_li_list = base_html.xpath("/html/body/div[3]/div[1]/ul/li")
# 保存所有文章的列表
news_text_dict_list = []
times = 1

# 数据库插入时的id
collection_insert_id = 1;

# 对新闻列表循环遍历
for news_li in news_li_list:
    # 获取新闻链接
    news_url = news_li.xpath("./h3/a/@href")[0]
    # 获取新闻标题
    news_title = news_li.xpath("./h3/a/text()")[0]
    news_source = news_li.xpath("./span/em[1]/text()")[0]
    news_date = news_li.xpath("./span/em[2]/text()")[0]
    # 继续爬取新闻的详情页
    news_response = requests.get(url=news_url, headers=headers)
    news_response.encoding = "utf-8"
    # 把详情页html信息赋值
    news_html = etree.HTML(news_response.text)

    news_text_div4_class = news_html.xpath("/html/body/div[3]/div[1]/div[4]/@class")
    # 有的页面有显示全文的按钮，因此要加入这个判断
    # 找到pages的div
    if news_text_div4_class is not None and len(news_text_div4_class) > 0 and news_text_div4_class[0] == "pages":
        news_text_div4_div_a_list = news_html.xpath("/html/body/div[3]/div[1]/div[4]/div/a")
        # 遍历div中的a标签
        for news_text_div4_div_a in news_text_div4_div_a_list:
            a_class = news_text_div4_div_a.xpath("./@class")
            # 找到nextPage的a标签，里面的href就是全文的新链接
            if a_class is not None and len(a_class) > 0 and a_class[0] == "nextPage":
                a_href = news_text_div4_div_a.xpath("./@href")[0]
                final_slash_index_in_url = news_url.rfind("/")
                news_url = news_url[:final_slash_index_in_url + 1] + a_href
                news_response.close()
                # 用更新后的全文的链接继续
                news_response = requests.get(url=news_url, headers=headers)
                # 乱码改编码
                news_response.encoding = "utf-8"
                # 把详情页html信息赋值
                news_html = etree.HTML(news_response.text)
                break

    # 段落标签的列表
    news_text_p_list = news_html.xpath("/html/body/div[3]/div[1]/div[3]/p")
    # 保存当前文章每个段落的字典
    news_text_dict = {"_id": collection_insert_id, "title": news_title, "source": news_source, "date": news_date}
    collection_insert_id += 1;
    index = 1;
    for news_text_p in news_text_p_list:
        news_text_p_text = news_text_p.xpath("./text()")
        # 根据网页结构，如果不是空的，说明是段落的文字
        if news_text_p_text is not None and len(news_text_p_text) > 0:
            # 保存段落文字到字典中
            news_text_dict[str(index)] = news_text_p_text[0]
            index += 1
        # 如果是空的，说明还有一个img的标签，是图片
        else:
            try:
                # 图片url
                news_text_p_img_src = news_text_p.xpath("./img/@src")[0]
                # 图片标题
                news_text_p_img_title = news_text_p.xpath("./img/@alt")[0]
                # 图片bytes数据
                news_text_p_img_bytes = requests.get(url=news_text_p_img_src, headers=headers).content
                # 保存图片标题和bytes数据到字典中
                # news_text_dict[str(index)] = [news_text_p_img_title, news_text_p_img_bytes]
                news_text_dict[str(index)] = [news_text_p_img_title, news_text_p_img_bytes]
                index += 1
            except:
                # 有粗体？
                # print(news_url)
                continue
    news_text_dict_list.append(news_text_dict)

    # 插入到数据库中
    collection_news.insert_one(news_text_dict)

    times += 1
    # if times == 2:
    #     break

    # 休眠1秒
    time.sleep(1)

# for news_text in news_text_dict_list:
#     print(news_text)

# 随机读取某一篇数据库中的新闻
# 写入到docx文档中进行展示
news_text_random_index = random.randint(1, len(news_text_dict_list))
news_text_dict = collection_news.find_one({"_id": news_text_random_index})

doc = Document()
doc.add_heading(news_text_dict["title"] + news_text_dict["source"] + news_text_dict["date"])
paragraph_index = 1;
while True:
    if str(paragraph_index) not in news_text_dict:
        break
    paragraph_content = news_text_dict[str(paragraph_index)]
    if isinstance(paragraph_content, list):
        image_title = paragraph_content[0]
        image_bytes = paragraph_content[1]
        with open("image.jpg", "wb") as fp:
            fp.write(image_bytes)
        doc.add_picture('image.jpg', width=shared.Cm(10))  # 按英寸设置
        doc.add_paragraph(image_title)
    else:
        doc.add_paragraph(paragraph_content)
    paragraph_index += 1
# 保存文件
doc.save('showNews.docx')

print("程序结束")
